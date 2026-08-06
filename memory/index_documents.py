# memory/index_documents.py — indexation des documents personnels
#
# ── À quoi ça sert ────────────────────────────────────────────────────
#
# Le RAG ne savait consulter que ce qu'on lui avait donné à la main, un
# fichier à la fois, depuis un interpréteur Python. Autant dire rien :
# la base ne contenait qu'un document d'exemple sur l'intelligence
# artificielle, qui polluait les réponses au lieu de les nourrir.
#
# Ce module lit un dossier et met la base à jour. Usage :
#
#     venv\Scripts\python.exe -m memory.index_documents
#     venv\Scripts\python.exe -m memory.index_documents --reset
#     venv\Scripts\python.exe -m memory.index_documents mes_docs/
#
# ── Ce qu'il garantit ─────────────────────────────────────────────────
#
# **Relancer la commande est sans danger.** Un fichier inchangé est
# ignoré (comparaison par empreinte du contenu), un fichier modifié voit
# ses anciens morceaux remplacés, et un fichier supprimé du disque est
# retiré de la base. Sans ça, chaque exécution empilait des doublons et
# la recherche remontait deux fois le même extrait.
#
# ── ⚠️ Sécurité ───────────────────────────────────────────────────────
#
# Ce sont les documents personnels de Cyril : contrats, relevés, notes.
# Rien ne sort de la machine — les embeddings sont calculés par Ollama en
# local (nomic-embed-text), et une question qui déclenche le RAG est
# forcée en local par route() (CLAUDE.md règle 3).
#
# Le CONTENU n'est jamais affiché ici, seulement les noms de fichiers et
# des compteurs : cette commande peut tourner devant quelqu'un.

from __future__ import annotations

import argparse
import re
import shutil
import sys
import tempfile
from pathlib import Path

if __package__ in (None, ""):  # exécution directe : python memory/index_documents.py
    sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from config import DOCUMENTS_DIR  # noqa: E402
from modules.ocr_engine import OCREngine, OCRUnavailable  # noqa: E402
from modules.rag_manager import RAGManager  # noqa: E402

# Formats lus sans dépendance supplémentaire.
#
# ⚠️ « .log » a été RETIRÉ le 01/08/2026. Un journal est volumineux,
# répétitif et sans valeur documentaire : celui de Cyril produisait 818
# morceaux sur 1083, soit 76 % de la base, et la recherche ne remontait
# pratiquement que lui — 35 vrais documents devenus introuvables.
TEXT_SUFFIXES = {".txt", ".md", ".markdown", ".rst", ".csv", ".json"}

# ⚠️ L'extension ne suffit PAS. Le fichier en cause s'appelait « log.txt » :
# retirer « .log » de la liste ci-dessus ne l'aurait pas écarté, et il
# serait revenu au passage suivant. D'où ce second test, sur le nom.
#
# Comparé au RADICAL (nom sans extension), pas en sous-chaîne : « log »
# dans « catalogue », « blog » ou « dialogue » ne doit rien déclencher.
LOG_SUFFIXES = {".log"}
LOG_STEMS = {"log", "logs", "journal", "debug", "trace", "output", "stdout", "stderr"}
LOG_PREFIXES = ("log_", "log-", "logs_", "logs-", "debug_", "debug-", "trace_")
LOG_SUFFIXES_IN_NAME = ("_log", "-log", "_logs", "-logs", "_debug", "-debug")

# Le PDF demande pypdf (installé le 01/08/2026 à la demande de Cyril :
# ses contrats et relevés sont dans ce format). Traité à part parce que
# la dépendance peut manquer sur une autre machine — le module doit
# rester utilisable sans elle.
PDF_SUFFIXES = {".pdf"}

# Le .docx demande python-docx (installé le 01/08/2026). Même logique que
# pypdf : traité à part, import paresseux, le module reste utilisable sans.
# ⚠️ Ne concerne PAS le .doc ancien, qui est un format binaire sans
# rapport — il reste dans KNOWN_UNSUPPORTED.
DOCX_SUFFIXES = {".docx"}

# Défini une fois : la liste était répétée à trois endroits, et « .pdf »
# avait déjà failli manquer dans l'un d'eux au moment de le brancher.
SUPPORTED_SUFFIXES = TEXT_SUFFIXES | PDF_SUFFIXES | DOCX_SUFFIXES

# ⚠️ Un PDF SCANNÉ ne contient aucune couche texte : pypdf en extrait une
# chaîne vide ou quelques caractères parasites. C'est le cas le plus
# courant pour un contrat reçu par la poste et photographié. En dessous
# de ce seuil, on considère qu'il n'y a rien à indexer et on le DIT —
# indexer trois caractères de bruit rendrait le document introuvable
# tout en le faisant apparaître comme traité.
PDF_MIN_CHARS = 80

# Le mode d'emploi du dossier n'est pas un document de Cyril. Indexé, il
# répondrait à « comment j'indexe mes documents ? » par lui-même, et
# polluerait les recherches comme le faisait sample_document.txt.
EXCLUDED_NAMES = {"readme.md", "readme.txt", "lisezmoi.txt"}

# ⚠️⚠️ REFUS DE SÉCURITÉ — fichiers de secrets ⚠️⚠️
#
# Le 01/08/2026, une indexation de C:/Users/PC/Documents a fait entrer en
# base « Mots de passe Microsoft Edge.csv » et « proton-recovery-kit.pdf ».
# Des identifiants en clair dans une base vectorielle deviennent
# RÉCUPÉRABLES par n'importe quelle question qui s'en approche
# sémantiquement — et le LLM les recopiera dans sa réponse sans état
# d'âme, puisqu'on les lui aura fournis comme contexte.
#
# Ce n'est pas du même ordre qu'un document personnel. Un bulletin de
# paie révèle un salaire ; un export de mots de passe donne l'accès aux
# comptes. D'où un REFUS, et non un simple avertissement.
#
# La liste porte sur le NOM du fichier, jamais sur son contenu : il
# faudrait l'ouvrir pour l'inspecter, ce qui est exactement ce qu'on
# cherche à éviter. Elle rate donc un fichier mal nommé — c'est un filet,
# pas une garantie. La garantie, c'est de ne pas ranger ses secrets dans
# un dossier qu'on indexe.
#
# Contournable avec --autoriser-secrets, jamais en silence.
# ⚠️ Cette liste a DÉJÀ raté un cas, le jour même où elle a été écrite :
# « Bitdefender SecurePass Recovery Key.pdf » n'a échappé à l'indexation
# que parce que c'était un scan sans couche texte. « recovery-kit » y
# était, « recovery key » non. C'est la démonstration que le filtre par
# nom est un filet avec des trous, et qu'il ne remplace pas de ranger
# ses secrets ailleurs que dans un dossier indexé.
SECRET_PATTERNS = (
    # mots de passe
    "mot de passe", "mots de passe", "motdepasse", "password", "passwd",
    "master password", "securepass", "keepass", "bitwarden", "lastpass",
    "1password", "dashlane", "nordpass", "vault", "coffre-fort",
    # récupération de compte
    "recovery", "recuperation", "récupération", "restauration",
    "codes de secours", "backup codes", "authenticator", "2fa", "otp",
    # clés et identifiants
    "identifiants", "credentials", "secret", "api-key", "api_key",
    "private-key", "private_key", "cle privee", "clé privée",
    "id_rsa", "id_ed25519", ".pem", "seed-phrase", "seed_phrase",
    "phrase de recuperation", "mnemonic", "wallet",
)

# Un fichier qui produit plus que cette part de la base l'écrase : la
# recherche ne remonterait pratiquement que lui. Mesuré : un log.txt de
# 0,7 Mo a produit 818 morceaux sur 1086, soit 75 % du total, noyant
# 38 vrais documents.
DOMINANCE_RATIO = 0.4

# Signalés explicitement plutôt qu'ignorés en silence : un document
# déposé qui n'apparaît jamais dans les réponses est un bug
# incompréhensible du point de vue de Cyril.
KNOWN_UNSUPPORTED = {
    ".doc": "format ancien — réenregistrer en .docx ou .txt",
    ".odt": "réenregistrer en .txt",
    ".xlsx": "exporter en .csv",
}


class UnreadablePDF(Exception):
    """PDF illisible — chiffré, corrompu, ou sans couche texte."""


def _extract_page(page) -> str:
    """
    Texte d'une page, en préservant l'association libellé → valeur.

    ⚠️ C'EST LE POINT DÉCISIF POUR LES DOCUMENTS EN COLONNES — bulletins
    de paie, factures, relevés. En extraction linéaire (le défaut),
    pypdf rendait :

        Cadre Net à payer NET A PAYER AVANT IMPOT SUR LE REVENU

    …sans aucun montant : les valeurs, situées dans une autre colonne,
    atterrissaient ailleurs dans le flux. Luca's recevait le bon document,
    le bon extrait, et ne pouvait toujours pas répondre.

    En mode « layout », la disposition spatiale est conservée :

        NET A PAYER AVANT IMPOT SUR LE REVENU | Cadre Net à payer | 1647.68
        MONTANT NET SOCIAL | 1625.68

    Les suites d'espaces sont réduites à « | » : le mode layout aligne
    par du remplissage (15 000 caractères contre 6 400), qui gonfle les
    morceaux sans rien apporter, et un séparateur explicite survit mieux
    au découpage et à l'embedding qu'une colonne d'espaces.
    """
    try:
        texte = page.extract_text(extraction_mode="layout") or ""
    except TypeError:
        # pypdf < 4 ne connaît pas extraction_mode. Mieux vaut un texte
        # dégradé que pas de texte du tout.
        return page.extract_text() or ""

    lignes = []
    for ligne in texte.splitlines():
        compacte = re.sub(r"[ \t]{2,}", " | ", ligne.strip())
        if compacte:
            lignes.append(compacte)
    return "\n".join(lignes)


def _rasteriser_pdf(path: Path) -> Path:
    """
    Rend chaque page du PDF en image PNG, dans un dossier temporaire.

    PyMuPDF (fitz) plutôt que pdf2image : pas de binaire externe
    (poppler) à installer, pour la même raison que RapidOCR a été choisi
    plutôt qu'easyocr dans modules/ocr_engine.py.

    Lève OCRUnavailable si la bibliothèque manque ou si le rendu échoue
    (PDF invalide, mot de passe, page corrompue) — l'appelant traite ça
    comme « OCR non tenté », pas comme « rien trouvé ».
    """
    try:
        import fitz  # PyMuPDF
    except ImportError as exc:
        raise OCRUnavailable("PyMuPDF absent — pip install pymupdf") from exc

    tmp_dir = Path(tempfile.mkdtemp(prefix="lucas_ocr_pdf_"))
    try:
        with fitz.open(str(path)) as document:
            for i, page in enumerate(document):
                pixmap = page.get_pixmap(dpi=200)
                pixmap.save(str(tmp_dir / f"page_{i:03d}.png"))
    except Exception as exc:  # noqa: BLE001 — fitz lève des types variés
        shutil.rmtree(tmp_dir, ignore_errors=True)
        raise OCRUnavailable(f"rendu PDF impossible ({type(exc).__name__})") from exc
    return tmp_dir


def _ocr_pdf(path: Path, ocr_engine: OCREngine) -> str:
    """
    Dernier recours pour un PDF sans couche texte : rendre chaque page en
    image, puis lire le texte par OCR (RapidOCR, en local — voir
    modules/ocr_engine.py).

    ⚠️ Les images intermédiaires sont supprimées dans tous les cas
    (succès ou échec) : ce sont potentiellement des pièces d'identité ou
    des contrats, elles ne doivent pas traîner dans le dossier temporaire
    du système au-delà de cette fonction.
    """
    images_dir = _rasteriser_pdf(path)
    try:
        morceaux = []
        for image_path in sorted(images_dir.iterdir()):
            resultat = ocr_engine.extract_text(image_path)
            if not resultat.is_empty:
                morceaux.append(resultat.text)
        return "\n\n".join(morceaux)
    finally:
        shutil.rmtree(images_dir, ignore_errors=True)


def _read_pdf(path: Path, ocr_engine: OCREngine | None = None) -> str:
    """
    Extrait le texte d'un PDF.

    ⚠️ Lève UnreadablePDF si ni la couche texte ni l'OCR (repli pour un
    PDF scanné — cartes d'identité, contrats reçus par la poste) ne
    donnent assez de matière. La distinction compte pour Cyril :
    « document vide, ignoré » ne lui dit pas quoi faire, « aucune couche
    texte, probablement scanné » lui dit que seul un OCR le rendrait
    consultable — et c'est justement ce qui vient d'être tenté.
    """
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise UnreadablePDF("pypdf absent — pip install pypdf") from exc

    try:
        reader = PdfReader(str(path))
        if reader.is_encrypted:
            # Un mot de passe vide suffit pour beaucoup de PDF « protégés
            # contre l'impression » — ça vaut la peine d'essayer avant
            # d'abandonner.
            try:
                reader.decrypt("")
            except Exception as exc:  # noqa: BLE001
                raise UnreadablePDF("PDF chiffré, mot de passe requis") from exc

        pages = [_extract_page(page) for page in reader.pages]
    except UnreadablePDF:
        raise
    except Exception as exc:  # noqa: BLE001 — pypdf lève des types variés
        raise UnreadablePDF(f"illisible ({type(exc).__name__})") from exc

    texte = "\n\n".join(p.strip() for p in pages if p.strip())
    if len(texte) >= PDF_MIN_CHARS:
        return texte

    try:
        texte_ocr = _ocr_pdf(path, ocr_engine or OCREngine())
    except OCRUnavailable:
        raise UnreadablePDF(
            f"aucune couche texte ({len(texte)} caractères sur "
            f"{len(reader.pages)} page(s)) — probablement scanné, "
            "OCR indisponible"
        ) from None

    if len(texte_ocr.strip()) < PDF_MIN_CHARS:
        raise UnreadablePDF(
            f"aucune couche texte — probablement scanné — et l'OCR n'a "
            f"rien trouvé d'exploitable sur {len(reader.pages)} page(s)"
        )
    return texte_ocr


class UnreadableDOCX(Exception):
    """.docx illisible — corrompu, protégé, ou sans contenu textuel."""


def _read_docx(path: Path) -> str:
    """
    Extrait le texte d'un .docx : paragraphes ET tableaux.

    ⚠️ Les tableaux comptent autant que les paragraphes. Un formulaire,
    une facture ou une demande d'autorisation d'absence mettent
    l'essentiel de leur information dedans — dates, montants, noms. Ne
    lire que `document.paragraphs`, comme le font la plupart des
    exemples, laisserait ces documents vides ou presque.
    """
    try:
        import docx
    except ImportError as exc:
        raise UnreadableDOCX("python-docx absent — pip install python-docx") from exc

    try:
        document = docx.Document(str(path))
    except Exception as exc:  # noqa: BLE001 — python-docx lève des types variés
        raise UnreadableDOCX(f"illisible ({type(exc).__name__})") from exc

    blocs = [p.text.strip() for p in document.paragraphs if p.text.strip()]

    for table in document.tables:
        for row in table.rows:
            cellules = [c.text.strip() for c in row.cells if c.text.strip()]
            # Les cellules fusionnées se répètent d'une colonne à l'autre :
            # sans déduplication, une ligne devient « Nom Nom Nom ».
            uniques = list(dict.fromkeys(cellules))
            if uniques:
                blocs.append(" | ".join(uniques))

    texte = "\n\n".join(blocs)
    if not texte.strip():
        raise UnreadableDOCX("aucun texte — document vide ou uniquement des images")
    return texte


def _read_text(path: Path) -> str | None:
    """
    Lit un fichier texte. Retourne None si l'encodage est illisible.

    Un seul fichier mal encodé ne doit pas interrompre l'indexation de
    tous les autres.
    """
    for encoding in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            return path.read_text(encoding=encoding)
        except (UnicodeDecodeError, LookupError):
            continue
        except OSError:
            return None
    return None


def _read(
    path: Path, ocr_engine: OCREngine | None = None
) -> tuple[str | None, str]:
    """
    Lit un document, quel que soit son format.

    Retourne (texte, motif) — `motif` explique l'échec quand `texte` est
    None, pour que la sortie dise POURQUOI un fichier n'est pas indexé.

    `ocr_engine` est réutilisé d'un appel à l'autre par `index_directory`
    (repli pour les PDF scannés) : recharger RapidOCR à chaque fichier
    coûterait cher sur un dossier qui en contient plusieurs.
    """
    suffix = path.suffix.lower()

    if suffix in PDF_SUFFIXES:
        try:
            return _read_pdf(path, ocr_engine), ""
        except UnreadablePDF as exc:
            return None, str(exc)

    if suffix in DOCX_SUFFIXES:
        try:
            return _read_docx(path), ""
        except UnreadableDOCX as exc:
            return None, str(exc)

    texte = _read_text(path)
    return (texte, "") if texte is not None else (None, "encodage illisible")


def looks_like_a_log(path: Path) -> bool:
    """
    Ce fichier est-il un journal ?

    Sur l'extension ET sur le radical du nom : le journal de Cyril
    s'appelait « log.txt », donc l'extension seule l'aurait laissé passer.
    """
    if path.suffix.lower() in LOG_SUFFIXES:
        return True

    stem = path.stem.lower().strip()
    return (
        stem in LOG_STEMS
        or stem.startswith(LOG_PREFIXES)
        or stem.endswith(LOG_SUFFIXES_IN_NAME)
    )


def looks_like_secrets(name: str) -> bool:
    """
    Ce nom de fichier annonce-t-il des identifiants ?

    Sur le NOM seul — inspecter le contenu supposerait de l'ouvrir et de
    le lire, ce qu'on cherche précisément à éviter pour ce genre de
    fichier.
    """
    normalise = name.lower().replace("_", " ").replace("-", " ")
    return any(
        motif.replace("_", " ").replace("-", " ") in normalise
        for motif in SECRET_PATTERNS
    )


def _collect(
    directory: Path, allow_secrets: bool = False
) -> tuple[list[Path], dict[str, list[str]], list[str], list[str]]:
    """Fichiers indexables, formats non gérés, secrets refusés, journaux."""
    indexable: list[Path] = []
    unsupported: dict[str, list[str]] = {}
    secrets: list[str] = []
    logs: list[str] = []

    for path in sorted(directory.rglob("*")):
        if not path.is_file() or path.name.startswith("."):
            continue
        if path.name.lower() in EXCLUDED_NAMES:
            continue

        # ⚠️ Le test « journal » passe avant le test d'extension : un
        # « log.txt » a une extension parfaitement lisible, c'est son nom
        # qui le disqualifie.
        if looks_like_a_log(path):
            logs.append(path.name)
            continue

        suffix = path.suffix.lower()
        if suffix not in SUPPORTED_SUFFIXES:
            if suffix in KNOWN_UNSUPPORTED:
                unsupported.setdefault(suffix, []).append(path.name)
            continue

        # ⚠️ Le refus passe AVANT la lecture : un fichier de secrets ne
        # doit pas être ouvert, même pour vérifier qu'il est lisible.
        if not allow_secrets and looks_like_secrets(path.name):
            secrets.append(path.name)
            continue

        indexable.append(path)

    return indexable, unsupported, secrets, logs


def _explain_unsupported(unsupported: dict[str, list[str]]) -> None:
    if not unsupported:
        return
    print("\nFormats non gérés pour l'instant :")
    for suffix, names in sorted(unsupported.items()):
        apercu = ", ".join(names[:3]) + (f" (+{len(names) - 3})" if len(names) > 3 else "")
        print(f"  {suffix:<7} {len(names):>3} fichier(s) — {KNOWN_UNSUPPORTED[suffix]}")
        print(f"          {apercu}")


def index_directory(
    directory: str | Path = DOCUMENTS_DIR,
    reset: bool = False,
    allow_secrets: bool = False,
) -> int:
    """
    Met la base à jour depuis un dossier. Retourne le code de sortie.
    """
    directory = Path(directory)
    if not directory.is_dir():
        print(f"Dossier introuvable : {directory}")
        print("Le créer et y déposer des documents, puis relancer.")
        return 1

    rag = RAGManager()
    if not rag.use_chroma or rag.collection is None:
        print("ChromaDB indisponible — indexation impossible.")
        print("Vérifier qu'Ollama tourne (ollama serve), il calcule les embeddings.")
        return 1

    if reset:
        connus = rag.indexed_documents()
        for doc_id in connus:
            rag.remove_document(doc_id)
        print(f"Base vidée : {len(connus)} document(s) retiré(s).\n")

    fichiers, non_geres, secrets, logs = _collect(
        directory, allow_secrets=allow_secrets
    )

    # ⚠️ Affiché AVANT l'indexation, et en tête : c'est ce que Cyril doit
    # voir même s'il ne lit que les premières lignes.
    if secrets:
        print("⚠️  REFUSÉ — fichiers de secrets, jamais indexés :")
        for nom in secrets:
            print(f"      {nom}")
        print(
            "    Des identifiants en base vectorielle deviennent récupérables\n"
            "    par une simple question, et le LLM les recopierait dans sa\n"
            "    réponse. Déplacer ces fichiers hors du dossier indexé.\n"
            "    Passer outre (à vos risques) : --autoriser-secrets\n"
        )

    if logs:
        apercu = ", ".join(logs[:4]) + (f" (+{len(logs) - 4})" if len(logs) > 4 else "")
        print(f"Journaux ignorés ({len(logs)}) : {apercu}\n")

    if not fichiers:
        print(f"Aucun fichier indexable dans {directory}")
        print(f"Formats lus : {', '.join(sorted(SUPPORTED_SUFFIXES))}")
        _explain_unsupported(non_geres)
        return 1

    print(f"{len(fichiers)} fichier(s) à examiner dans {directory}\n")

    ajoutes = inchanges = illisibles = 0
    vus: set[str] = set()
    morceaux_par_doc: dict[str, int] = {}
    # Une seule instance pour tout le dossier : recharger RapidOCR à
    # chaque PDF scanné coûterait cher si plusieurs en contiennent.
    ocr_engine = OCREngine()

    for path in fichiers:
        doc_id = path.name
        vus.add(doc_id)

        texte, motif = _read(path, ocr_engine)
        if texte is None:
            print(f"  illisible  {doc_id} — {motif}")
            illisibles += 1
            continue

        # add_text() rend False quand le contenu est identique à ce qui
        # est déjà en base — c'est ce qui rend la commande relançable.
        if rag.add_text(texte, doc_id):
            ajoutes += 1
        else:
            print(f"  inchangé   {doc_id}")
            inchanges += 1
        morceaux_par_doc[doc_id] = len(rag._chunk_text(texte))

    # Un fichier retiré du disque resterait consultable indéfiniment : la
    # base n'a aucun moyen de l'apprendre autrement qu'ici.
    orphelins = rag.indexed_documents() - vus
    for doc_id in sorted(orphelins):
        morceaux = rag.remove_document(doc_id)
        print(f"  retiré     {doc_id} ({morceaux} morceaux — absent du dossier)")

    total = len(rag.collection.get(include=[])["ids"])

    # Un fichier qui écrase la base rend tous les autres introuvables. Un
    # log.txt de 0,7 Mo a produit 818 morceaux sur 1086 — 75 % du total,
    # noyant 38 vrais documents. Signalé, pas retiré d'office : c'est
    # peut-être un document légitimement volumineux.
    for doc_id, morceaux in sorted(morceaux_par_doc.items(), key=lambda x: -x[1]):
        if total and morceaux / total > DOMINANCE_RATIO:
            print(
                f"\n⚠️  {doc_id} pèse {morceaux} morceaux sur {total} "
                f"({morceaux / total:.0%} de la base).\n"
                "    Il écrase les autres documents : la recherche ne remontera\n"
                "    pratiquement que lui. Le retirer du dossier indexé, ou :\n"
                f"    RAGManager().remove_document({doc_id!r})"
            )
        break

    print(
        f"\n{ajoutes} indexé(s), {inchanges} inchangé(s), "
        f"{len(orphelins)} retiré(s), {illisibles} illisible(s)"
        f"\nbase : {total} morceaux"
    )
    _explain_unsupported(non_geres)

    # Le seuil de pertinence est calibré sur DEUX morceaux d'un document
    # d'exemple, avec une marge étroite (voir RAG_MAX_DISTANCE). Il n'a
    # plus aucune raison d'être juste maintenant que la base a changé.
    if ajoutes or orphelins:
        print(
            "\n➜ La base a changé : recalibrer le seuil de pertinence.\n"
            "  venv\\Scripts\\python.exe demos\\calibrate_rag.py"
        )
    return 0


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(
        description="Indexe les documents personnels pour le RAG (tout reste local).",
    )
    parser.add_argument(
        "directory", nargs="?", default=DOCUMENTS_DIR,
        help=f"dossier à indexer (défaut : {DOCUMENTS_DIR})",
    )
    parser.add_argument(
        "--reset", action="store_true",
        help="vide la base avant d'indexer (utile pour retirer les documents d'exemple)",
    )
    parser.add_argument(
        "--autoriser-secrets", dest="allow_secrets", action="store_true",
        help="indexe AUSSI les fichiers qui ressemblent à des mots de passe "
             "ou des clés de récupération (fortement déconseillé)",
    )
    args = parser.parse_args(argv)
    return index_directory(
        args.directory, reset=args.reset, allow_secrets=args.allow_secrets
    )


if __name__ == "__main__":
    raise SystemExit(main())
